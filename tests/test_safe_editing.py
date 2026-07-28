from pathlib import Path

import fitz
import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from pocket_agent.config.models import PathsConfig
from pocket_agent.tools.files.docx_edit import modify_docx
from pocket_agent.tools.files.excel import modify_excel
from pocket_agent.tools.files.pdf_edit import modify_pdf
from pocket_agent.tools.files.safety import SafeEditSession, file_checksum, prepare_safe_edit


def _paths(tmp_path: Path) -> PathsConfig:
    nas = tmp_path / "nas"
    nas.mkdir()
    return PathsConfig(
        {
            "nas": {"root": str(nas), "allowed_read_roots": [str(nas)]},
            "data": {
                "logs": "data/logs",
                "working": "data/working",
                "backup": "data/backup",
            },
        },
        tmp_path,
    )


def _make_xlsx(nas: Path) -> Path:
    path = nas / "budget.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Budget"
    ws["A1"] = "Item"
    ws["B1"] = "Amount"
    ws["A2"] = "Rent"
    ws["B2"] = 1000
    wb.save(path)
    return path


def _make_docx(nas: Path) -> Path:
    path = nas / "notes.docx"
    doc = Document()
    doc.add_paragraph("line one")
    doc.save(path)
    return path


def _make_pdf(nas: Path) -> Path:
    path = nas / "doc.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "original")
    doc.save(path)
    doc.close()
    return path


@pytest.mark.asyncio
async def test_prepare_safe_edit_creates_backup_and_working(tmp_path: Path):
    paths = _paths(tmp_path)
    paths.logs_dir.mkdir(parents=True, exist_ok=True)
    nas_file = _make_xlsx(paths.nas_root)
    original_checksum = file_checksum(nas_file)

    session = prepare_safe_edit(paths, nas_file, "test")
    assert isinstance(session, SafeEditSession)
    assert session.backup_path.is_file()
    assert session.working_path.is_file()
    assert session.original_checksum == original_checksum


@pytest.mark.asyncio
async def test_modify_excel_cell(tmp_path: Path):
    paths = _paths(tmp_path)
    paths.logs_dir.mkdir(parents=True, exist_ok=True)
    xlsx = _make_xlsx(paths.nas_root)

    result = await modify_excel(paths, str(xlsx), "Budget", "B2", "2500")
    assert result.success

    wb = load_workbook(xlsx, data_only=True)
    assert wb["Budget"]["B2"].value == 2500
    wb.close()

    assert Path(result.data["backup_path"]).is_file()


@pytest.mark.asyncio
async def test_modify_docx_append(tmp_path: Path):
    paths = _paths(tmp_path)
    paths.logs_dir.mkdir(parents=True, exist_ok=True)
    docx = _make_docx(paths.nas_root)

    result = await modify_docx(paths, str(docx), "new line", action="append")
    assert result.success

    doc = Document(docx)
    assert len(doc.paragraphs) == 2
    assert doc.paragraphs[-1].text == "new line"


@pytest.mark.asyncio
async def test_modify_pdf_add_page(tmp_path: Path):
    paths = _paths(tmp_path)
    paths.logs_dir.mkdir(parents=True, exist_ok=True)
    pdf = _make_pdf(paths.nas_root)

    result = await modify_pdf(paths, str(pdf), "added page text", action="add_page")
    assert result.success

    with fitz.open(pdf) as doc:
        assert doc.page_count == 2
