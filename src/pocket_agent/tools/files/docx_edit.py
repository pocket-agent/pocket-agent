from docx import Document

from pocket_agent.config.models import PathsConfig
from pocket_agent.logging.action_log import log_action
from pocket_agent.tools.base import ToolResult
from pocket_agent.tools.files.paths import resolve_read_path
from pocket_agent.tools.files.safety import finalize_safe_edit, prepare_safe_edit
from pocket_agent.tools.files.validators import validate_docx


async def modify_docx(
    paths: PathsConfig,
    file_path: str,
    text: str,
    action: str = "append",
) -> ToolResult:
    """Append or replace-last paragraph with safe edit pipeline."""
    if action not in {"append", "replace_last"}:
        return ToolResult(success=False, error=f"Unsupported Word action: {action}")

    resolved = resolve_read_path(paths, file_path)
    if isinstance(resolved, ToolResult):
        log_action(
            paths.logs_dir,
            "modify_docx",
            {"file_path": file_path},
            success=False,
            error=resolved.error,
        )
        return resolved

    path = resolved
    if path.suffix.lower() != ".docx":
        return ToolResult(success=False, error="File is not a Word document (.docx)")

    session = prepare_safe_edit(paths, path, "modify_docx")
    if isinstance(session, ToolResult):
        return session

    try:
        doc = Document(session.working_path)
        if action == "append":
            doc.add_paragraph(text)
        else:
            paragraphs = doc.paragraphs
            if not paragraphs:
                doc.add_paragraph(text)
            else:
                paragraphs[-1].text = text
        doc.save(session.working_path)
    except Exception as exc:
        log_action(
            paths.logs_dir,
            "modify_docx",
            {"file_path": str(path)},
            success=False,
            error=str(exc),
        )
        return ToolResult(success=False, error=str(exc))

    result = finalize_safe_edit(paths, session, validate_docx, "modify_docx")
    if result.success:
        result.data["action"] = action
        result.data["text_length"] = len(text)
    return result
